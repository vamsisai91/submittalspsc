import streamlit as st
import fitz  # PyMuPDF
import re
from openpyxl import Workbook
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openai
from dotenv import load_dotenv
import os
import faiss
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter
import requests
import time
from openai import Client
from assistant import run_OpenAI_assistant
from pypdf import PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

# Load environment variables
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

# Function to extract text from specified table of content pages of uploaded PDF (Specs)
def extract_text_from_pdf(file, start_page, end_page):
    document = fitz.open(stream=file, filetype="pdf")
    text = ""
    for page_num in range(start_page - 1, end_page):
        page = document.load_page(page_num)
        text += page.get_text()
    return text

# Function to extract text from entire PDF
def extract_full_text_from_pdf(file):
    document = fitz.open(stream=file, filetype="pdf")
    text = ""
    for page_num in range(document.page_count):
        page = document.load_page(page_num)
        text += page.get_text()
    return text

# Function to extract unique section numbers from the text
def extract_section_numbers(text, section_pattern=None):
    if section_pattern is None:
        section_pattern = re.compile(
            r'(\b\d{2} \d{2} \d{2}\b|\b\d{6}\b|\b\d{3} \d{3}\b|\b\d{2} \d{4}\b|'
            r'\b\d{5}\b|\b\d{2} \d{3}\b|\b\d{3} \d{2}\b|'
            r'\b\d{4}\b|\b\d{2} \d{2}\b|'
            r'\b\d{3} \-)', re.MULTILINE)  
    section_numbers = section_pattern.findall(text)
    seen = set()
    unique_section_numbers = [x for x in section_numbers if not (x in seen or seen.add(x))]
    return unique_section_numbers

# Function to find addons for section numbers
def find_addons(text, section_numbers):
    addons = []
    for section in section_numbers:
        addon_pattern = re.compile(rf'{re.escape(section)}\.\d{{2}}')
        addons.extend(addon_pattern.findall(text))
    return addons

# Function to extract specific section based on heading and capture the section name
def extract_section(text, section_heading):
    pattern = re.compile(rf'({section_heading}\s+.*?END OF SECTION)', re.DOTALL)
    match = pattern.search(text)
    if match:
        name_pattern = re.compile(rf'{section_heading}\s+(.*?)\n')
        name_match = name_pattern.search(match.group(1))
        section_name = name_match.group(1) if name_match else 'Unknown'
        return match.group(1), section_name
    return None, None

# Function to extract submittals subsection
def extract_submittals_subsection(text):
    submittal_types = ["SUBMITTALS", "ACTION SUBMITTALS", "INFORMATIONAL SUBMITTALS", "CLOSEOUT SUBMITTALS", "SHOP DRAWING SUBMITTALS"]
    submittals = []
    for submittal_type in submittal_types:
        pattern = re.compile(rf'({submittal_type}.*?)(?=\n\d+\.\d+|\Z)', re.DOTALL)
        match = pattern.search(text)
        if match:
            submittals.append(match.group(1))
    return "\n\n".join(submittals) if submittals else None

# Function to sanitize sheet titles
def sanitize_sheet_title(title):
    invalid_chars = ['/', '\\', '?', '*', '[', ']']
    for char in invalid_chars:
        title = title.replace(char, '')
    return title[:31]  # Limit to 31 characters

# Function to add a new heading with a page break
def add_heading_with_page_break(doc, heading_text):
    doc.add_page_break()
    heading = doc.add_heading(level=1)
    run = heading.add_run(heading_text)
    run.bold = True
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add this function to create PDF
def create_pdf(project_name, all_extracted_content):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Add project name and title
    story.append(Paragraph(project_name, styles['Title']))
    story.append(Paragraph("EXTRACTED SUBMITTALS", styles['Title']))
    story.append(PageBreak())

    # Use existing styles or create new ones if they don't exist
    if 'Heading1' not in styles:
        styles.add(ParagraphStyle(name='Heading1', fontSize=14, spaceAfter=12))
    if 'BodyText' not in styles:
        styles.add(ParagraphStyle(name='BodyText', fontSize=10, spaceAfter=6))

    # Process content
    sections = all_extracted_content.strip().split('\n\n')
    for section in sections:
        section_lines = section.split('\n')
        heading_text = section_lines[0]
        content = section_lines[1:]

        story.append(Paragraph(heading_text, styles['Heading1']))
        for line in content:
            story.append(Paragraph(line, styles['BodyText']))
        story.append(PageBreak())

    doc.build(story)
    buffer.seek(0)
    return buffer

# Function to chunk text
def chunk_text(text):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    chunks = text_splitter.split_text(text)
    return chunks

client = Client()
# Function to get embeddings from OpenAI
def get_embeddings(text_list):
    embeddings = []
    for text in text_list:
        response = client.embeddings.create(
            input=[text], 
            model="text-embedding-ada-002"
        )
        embeddings.append(response.data[0].embedding)
    return np.array(embeddings)

# Function to store chunks and embeddings in FAISS
def store_embeddings_in_faiss(chunks, embeddings):
    dimension = len(embeddings[0])
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings)
    return index, chunks

# Function to query FAISS index and get relevant chunks
def query_faiss_index(query, index, chunks):
    query_embedding = get_embeddings([query])[0]
    D, I = index.search(np.array([query_embedding]), k=5)  # Get top 5 relevant chunks
    return [chunks[i] for i in I[0]]

# Function to get a response from OpenAI based on relevant text chunks
def get_openai_response(query, relevant_chunks):
    context = "\n\n".join(relevant_chunks)
    prompt = f"Context: {context}\n\nQuery: {query}\n\nResponse:"
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an assistant that will extract information from a user uploaded pdf"},
            {"role": "user", "content": prompt},
        ],
        max_tokens=500
    )
    return response.choices[0].message.content.strip()

# Streamlit UI Layout
st.set_page_config(layout="wide")

# Create a container for the header
header = st.container()

# Inside the container, create three columns
left_col, middle_col, right_col = header.columns([2, 1, 1])

# Add the title to the left column
with left_col:
    st.markdown(
        """
        <h1 style='font-size: 48px; margin-bottom: 0; white-space: nowrap;'>PS Submittals Extraction</h1>
        """,
        unsafe_allow_html=True
    )

# Add the video tutorial link to the right column
with right_col:
    st.markdown(
        """
        <a href="https://youtu.be/AQSR5KJ51SY" target="_blank" style="text-decoration: none; color: inherit;">
            <div style="float: right; text-align: center;">
                <img src="https://img.icons8.com/color/48/000000/youtube-play.png" alt="Video Tutorial">
                <br>
                <span style="font-size: 12px;">Video Tutorial</span>
            </div>
        </a>
        """,
        unsafe_allow_html=True
    )

st.image("logo.png", width=300)
st.markdown("<h2 style='text-align: left;'>We Build Life Changing Infrastructure</h2>", unsafe_allow_html=True)

# Define session state keys
if 'section_numbers_array' not in st.session_state:
    st.session_state.section_numbers_array = None
if 'pdf_file' not in st.session_state:
    st.session_state.pdf_file = None
if 'project_name' not in st.session_state:
    st.session_state.project_name = None
if 'special_section_number' not in st.session_state:
    st.session_state.special_section_number = None
if 'output_excel_path' not in st.session_state:
    st.session_state.output_excel_path = None
if 'output_path' not in st.session_state:
    st.session_state.output_path = None
if 'all_extracted_content' not in st.session_state:
    st.session_state.all_extracted_content = ""

st.header("Upload PDF and Provide Inputs about the Project and its Table of Contents (TOC)")
uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
project_name = st.text_input("Enter the name of the project:")
start_page = st.number_input("Enter the starting page number of the TOC:", min_value=1)
end_page = st.number_input("Enter the ending page number of the TOC:", min_value=start_page)
special_section_number = st.text_input("Enter the Submittals Master Section number (e.g., '01300' OR '01 33 00'):")

if st.button("Extract Section Numbers"):
    if uploaded_file is not None and start_page <= end_page and project_name and special_section_number:
        pdf_file = uploaded_file.read()
        pdf_text = extract_text_from_pdf(pdf_file, start_page, end_page)
        section_numbers = extract_section_numbers(pdf_text)
        
        if section_numbers:
            st.write("Extracted Section Numbers:")
            st.write(section_numbers)
            
            addons = find_addons(pdf_text, section_numbers)
            all_section_numbers = section_numbers + [addon for addon in addons if addon not in section_numbers]
            
            st.write("All Section Numbers (including addons):")
            st.write(all_section_numbers)
            
            st.session_state.section_numbers_array = all_section_numbers
            st.session_state.pdf_file = pdf_file
            st.session_state.project_name = project_name
            st.session_state.special_section_number = special_section_number

if st.session_state.section_numbers_array and st.session_state.pdf_file:
    if st.button("Confirm and Extract Documents"):
        pdf_file = st.session_state.pdf_file
        pdf_text = extract_full_text_from_pdf(pdf_file)
        section_numbers_array = st.session_state.section_numbers_array
        special_section_number = st.session_state.special_section_number
        project_name = st.session_state.project_name

        # Initialize a variable to hold all extracted sections and submittals
        all_extracted_content = ""

        # Handle the special section separately to extract the entire section
        special_section_heading = f"SECTION {special_section_number}"
        special_section, special_section_name = extract_section(pdf_text, special_section_heading)

        if special_section:
            all_extracted_content += f"{special_section_heading} - {special_section_name}\n{special_section}\n\n"

        # Remove the special section from the section numbers array
        section_numbers_array = [number for number in section_numbers_array if number != special_section_number]

        # Iterate over each section number, extract the section and then extract the "SUBMITTALS" subsection
        toc_entries = []
        for section_number in section_numbers_array:
            section_heading = f"SECTION {section_number}"
            extracted_section, section_name = extract_section(pdf_text, section_heading)

            if extracted_section:
                submittals_subsection = extract_submittals_subsection(extracted_section)

                if submittals_subsection:
                    all_extracted_content += f"{section_heading} - {section_name}\n{submittals_subsection}\n\n"
                    toc_entries.append(f"{section_heading} - {section_name}")

        # Create an Excel workbook
        wb = Workbook()

        # Add the project name and title on the first sheet
        ws = wb.active
        ws.title = sanitize_sheet_title("Project Info")
        ws.append([project_name])
        ws.append(["EXTRACTED SUBMITTALS"])

        # Write each section's content to separate sheets
        sections = all_extracted_content.strip().split('\n\n')
        for section in sections:
            section_lines = section.split('\n')
            heading_text = section_lines[0]
            content = section_lines[1:]

            # Sanitize and truncate the sheet title
            sheet_title = sanitize_sheet_title(heading_text)
            ws = wb.create_sheet(title=sheet_title)

            # Store the full section title in the first cell
            ws.append([heading_text])
            for line in content:
                ws.append([line])

        # Save the Excel workbook with the project name in the title
        output_excel_path = f'{project_name}_Extracted_SUBMITTALS_Sections.xlsx'
        wb.save(output_excel_path)
        st.session_state.output_excel_path = output_excel_path
        st.write(f"All sections and 'SUBMITTALS' subsections extracted and saved to {output_excel_path}")

        # Create a Word document
        doc = Document()

        # Add the project name and title on the first page
        project_title = doc.add_heading(project_name, level=1)
        project_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        extracted_submittals_title = doc.add_heading('EXTRACTED SUBMITTALS', level=1)
        extracted_submittals_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Create table of contents (TOC)
        toc = doc.add_paragraph()
        run = toc.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = r'TOC \o "1-3" \h \z \u'
        run._r.append(instrText)
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar)
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

        # Add each section and its submittals to the document, starting each section on a new page
        for section in all_extracted_content.strip().split('\n\n'):
            section_lines = section.split('\n')
            heading_text = section_lines[0]
            content = '\n'.join(section_lines[1:])

            add_heading_with_page_break(doc, heading_text)
            doc.add_paragraph(content)

        # Save the document with the project name in the title
        output_path = f'{project_name}_Extracted_SUBMITTALS_Sections.docx'
        doc.save(output_path)
        st.session_state.output_path = output_path
        st.write(f"All sections and 'SUBMITTALS' subsections extracted and saved to {output_path}")

        # Create a Word document for the TOC
        toc_doc = Document()

        # Add TOC title
        toc_title = toc_doc.add_heading('Table of Contents', level=1)
        toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add TOC entries to the TOC document
        for entry in toc_entries:
            toc_entry = toc_doc.add_paragraph()
            toc_entry.add_run(entry)

        # Save the TOC document
        toc_output_path = f'{project_name}_TOC.docx'
        toc_doc.save(toc_output_path)
        st.session_state.toc_output_path = toc_output_path
        st.write(f"Table of Contents extracted and saved to {toc_output_path}")
        
        st.session_state.all_extracted_content = all_extracted_content

        # Create PDF
        pdf_buffer = create_pdf(st.session_state.project_name, st.session_state.all_extracted_content)
        pdf_output_path = f'{st.session_state.project_name}_Extracted_SUBMITTALS_Sections.pdf'
        with open(pdf_output_path, 'wb') as f:
            f.write(pdf_buffer.getvalue())
        st.session_state.pdf_output_path = pdf_output_path
        st.write(f"All sections and 'SUBMITTALS' subsections extracted and saved to {pdf_output_path}")

# Display download buttons if documents are generated
if st.session_state.output_excel_path and st.session_state.output_path:
    with open(st.session_state.output_excel_path, "rb") as file:
        st.download_button(
            label="Download all the Submittals in an Excel File",
            data=file,
            file_name=st.session_state.output_excel_path,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    with open(st.session_state.output_path, "rb") as file:
        st.download_button(
            label="Download all the Submittals in a DOCX File",
            data=file,
            file_name=st.session_state.output_path,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    if 'toc_output_path' in st.session_state:
        with open(st.session_state.toc_output_path, "rb") as file:
            st.download_button(
                label="Download all the Submittals TOC (Submittal Schedule) in a DOCX File",
                data=file,
                file_name=st.session_state.toc_output_path,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
if 'pdf_output_path' in st.session_state:
    with open(st.session_state.pdf_output_path, "rb") as file:
        st.download_button(
            label="Download all the Submittals in a PDF File (Chat with OpenAI)",
            data=file,
            file_name=st.session_state.pdf_output_path,
                mime="application/pdf"
            )
# Add new sections for comparison

col1, col2 = st.columns(2)

with col1:
    st.header("Chat with JEA's Vol3 Approved Materials list")
    user_input_manual = st.text_input("You: ", key="user_input_manual")
    if st.button("Chat", key="send_manual"):
        if user_input_manual:
            assistant_id = "asst_EZQ9NL71x9QXNrncnzTqZMWv"  # Your assistant ID
            try:
                # Call the run_OpenAI_assistant function to get the response
                assistant_response = run_OpenAI_assistant(
                    assistant_id=assistant_id,
                    prompt=user_input_manual,
                    model='gpt-4o-mini'
                )
                
                # Extract the main text content (assuming it's the first item in the list)
                response_text = assistant_response[0] if isinstance(assistant_response, list) else assistant_response
                
                # Display the response with custom styling
                st.markdown(
                    f"""
                    <div style='background-color: #ffffff; color: #000000; padding: 10px; border-radius: 10px; overflow-wrap: break-word;'>
                        <p style='color: #000000; font-size: 16px; line-height: 1.5;'>{response_text}</p>
                    </div>
                    """, 
                    unsafe_allow_html=True
                )
            except Exception as e:
                st.write("Error: ", str(e))

with col2:
    st.header("Upload the Extracted Submittal/Project's Specifications/Standard Manual's here")
    uploaded_specifications = st.file_uploader("Choose a PDF file", type="pdf", key="specifications")
    
    if uploaded_specifications:
        # Compute embeddings only if they haven't been computed yet
        if 'specifications_embeddings' not in st.session_state:
            st.session_state.specifications_text = extract_full_text_from_pdf(uploaded_specifications.read())
            specifications_chunks = chunk_text(st.session_state.specifications_text)
            specifications_embeddings = get_embeddings(specifications_chunks)
            st.session_state.specifications_index, st.session_state.specifications_stored_chunks = store_embeddings_in_faiss(specifications_chunks, specifications_embeddings)
            st.session_state.specifications_embeddings = specifications_embeddings  # Store embeddings to prevent recomputation
            st.success("Specifications uploaded and embeddings computed successfully.")
        else:
            st.success("Specifications embeddings are already computed.")
    
    st.header("Chat with your Uploaded Document")
    user_input_specifications = st.text_input("You: ", key="user_input_specifications")
    if st.button("Chat", key="send_specifications"):
        if user_input_specifications and 'specifications_index' in st.session_state:
            # Query FAISS index using the stored embeddings and index
            relevant_chunks_specifications = query_faiss_index(
                user_input_specifications,
                st.session_state.specifications_index,
                st.session_state.specifications_stored_chunks
            )
            
            # Get OpenAI response based on relevant chunks
            response_specifications = get_openai_response(user_input_specifications, relevant_chunks_specifications)
            st.write("OpenAI: ", response_specifications)
        else:
            st.warning("Please upload a pdf document first.")


# Add footer
st.markdown(
    """
    <style>
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: white;
        color: black;
        text-align: center;
        padding: 10px;
    }
    </style>
    <div class="footer">
        <p>© 2024 Petticoat Schmitt. All rights reserved.</p>
    </div>
    """,
    unsafe_allow_html=True
)
